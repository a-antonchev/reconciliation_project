import os
import pathlib

import pandas as pd
import docx


def parse_excel_to_markdown(file_path: str) -> str:
    """
    Читает Excel файл (все листы) и конвертирует их в Markdown таблицы.
    """

    try:
        # читаем все листы, игнорируем заголовок
        xls = pd.read_excel(file_path, sheet_name=None, header=None)

        markdown_content = []

        # по каждому листу, датафрейму:
        for sheet_name, df in xls.items():
            markdown_content.append(f"#### Лист Excel: {sheet_name}\n")

            # удаляем полностью пустые строки и столбцы
            df = df.dropna(axis=0, how="all")
            df = df.dropna(axis=1, how="all")

            # заполняем пропустки (NaN) пустыми строками
            df = df.fillna("")

            if not df.empty:
                # конвертируем датафрейм в `markdown`
                markdown_content.append(df.to_markdown(index=False))

            # после каждого листа добавляем разделитель `\n\n`
            markdown_content.append("\n\n")

        # объединяем элементы списка через пустую строку
        return "".join(markdown_content)

    except Exception as e:
        raise RuntimeError(f"Ошибка при чтении Excel файла {file_path}: {e}!")


def parse_docx_to_markdown(file_path: str) -> str:
    """
    Читает Word файл, извлекает текст и конвертирует таблицы в Markdown.
    """

    try:
        doc = docx.Document(file_path)

        markdown_content = []

        # извлекаем обычный тескт
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                markdown_content.append(text)

        markdown_content.append("\n### Таблицы из документа:\n")

        # извлекаем таблицы и перерисовываем их в markdown
        for table in doc.tables:
            for i, row in enumerate(table.rows):

                # звлекаем текст из ячеек, заменяя переносы строк на пробелы, 
                # иначе markdown-таблица сломается
                row_data = [
                    cell.text.strip().replace("\n", " ") for cell in row.cells
                ]

                # склеиваем элементы через разделители `|` и рисуем markdown-строку таблицы
                markdown_content.append("| " + " | ".join(row_data) + " |")
                
                # если это первая строка, то добавляем разделитель по количеству ячеек (row.cells) в строке
                if i == 0:
                    markdown_content.append("|" + "---|" * len(row.cells))

            # после каждой таблицы добавляем разделитель `\n`
            markdown_content.append("\n")

        # склеиваем элементы через `\n`
        return "\n".join(markdown_content)

    except Exception as e:
        raise RuntimeError(f"Ошибка при чтении Word файла {file_path}: {e}!")


def parse_file(file_path: str) -> str:
    _, ext = os.path.splitext(file_path)

    # TODO: переписать через dict:
    
    if ext in [".xls", ".xlsx"]:
        return parse_excel_to_markdown(file_path)

    elif ext == ".docx":
        return parse_docx_to_markdown(file_path)

    elif ext == ".doc":
        raise ValueError(
            "Старый формат Word (.doc) не поддерживается. "
            "Пожалуйста, откройте файл и пересохраните его в современном формате (.docx)."
        )

    else:
        raise ValueError(
            f"Неподдерживаемый формат файла: {ext}. Разрешены только .xlsx, .xls и .docx"
        )


if __name__ == "__main__":
    data_dir = "./data"
    test_file = "specification.xlsx"

    # file_path = pathlib.Path(data_dir) / test_file
    file_path = os.path.join(data_dir, test_file)

    if pathlib.Path(file_path).is_file():
        markdown_result = parse_file(file_path)

        output_file = "specification.md"
        pathlib.Path(output_file).write_text(markdown_result, encoding="utf-8")
        # print(markdown_result)
    else:
        print(f"Создайте файл {test_file} в каталоге {data_dir}.")
